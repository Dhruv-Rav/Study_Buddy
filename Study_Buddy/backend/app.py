# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from transformers import pipeline
# import tensorflow as tf
# from tensorflow import keras
# import logging

# app = Flask(__name__)
# CORS(app)

# # Initialize the summarization pipeline
# summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# # Configure logging
# logging.basicConfig(level=logging.DEBUG)

# @app.route('/summarize', methods=['POST'])
# def summarize_text():
#     data = request.json
#     text = data.get('text', '')
    
#     # Log the received text
#     app.logger.debug(f"Received text: {text}")
    
#     # Perform summarization
#     try:
#         summary = summarizer(text, max_length=80, min_length=30, do_sample=False)
#         summary_text = summary[0]['summary_text']
        
#         # Log the summary
#         app.logger.debug(f"Summary: {summary_text}")
        
#         return jsonify({'summary': summary_text})
#     except Exception as e:
#         app.logger.error(f"Error during summarization: {e}")
#         return jsonify({'summary': 'An error occurred during summarization.'}), 500

# if __name__ == '__main__':
#     app.run(debug=True)

from flask import Flask, request, jsonify
from flask_cors import CORS
from transformers import pipeline
import logging
import os
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import docx  # Correct import for python-docx

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})  # Allow all origins

# Initialize the summarization and question-answering pipelines
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
qa_pipeline = pipeline("question-answering")

# Configure logging
logging.basicConfig(level=logging.DEBUG)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    text = ""
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        elif ext == 'pdf':
            with open(filepath, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
        elif ext == 'docx':
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text
        elif ext == 'doc':
            # Handle .doc files if needed
            pass
    except Exception as e:
        app.logger.error(f"Error extracting text from file: {e}")
    return text

def split_text(text, max_length=512):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        if len(current_chunk) + len(word.split()) <= max_length:
            current_chunk.append(word)
        else:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

@app.route('/summarize', methods=['POST'])
def summarize_text():
    if 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            text = extract_text_from_file(filepath)
            if not text:
                return jsonify({'error': 'Failed to extract text from file'}), 500
            try:
                app.logger.debug(f"Extracted text: {text[:500]}...")  # Log the first 500 characters of the extracted text
                chunks = split_text(text)
                summaries = [summarizer(chunk, max_length=100, min_length=50, do_sample=False)[0]['summary_text'] for chunk in chunks]
                summary_text = ' '.join(summaries)
                app.logger.debug(f"Summary: {summary_text}")
                return jsonify({'summary': summary_text})
            except Exception as e:
                app.logger.error(f"Error during summarization: {e}")
                return jsonify({'summary': 'An error occurred during summarization.'}), 500
        return jsonify({'error': 'Invalid file type'}), 400
    else:
        data = request.json
        text = data.get('text', '')
        
        # Log the received text
        app.logger.debug(f"Received text: {text}")
        
        # Perform summarization
        try:
            chunks = split_text(text)
            summaries = [summarizer(chunk, max_length=100, min_length=50, do_sample=False)[0]['summary_text'] for chunk in chunks]
            summary_text = ' '.join(summaries)
            
            # Log the summary
            app.logger.debug(f"Summary: {summary_text}")
            
            return jsonify({'summary': summary_text})
        except Exception as e:
            app.logger.error(f"Error during summarization: {e}")
            return jsonify({'summary': 'An error occurred during summarization.'}), 500

@app.route('/question_answer', methods=['POST'])
def question_answer():
    data = request.json
    context = data.get('context', '')
    question = data.get('question', '')
    
    # Log the received question and context
    app.logger.debug(f"Received question: {question}")
    app.logger.debug(f"Context: {context[:500]}...")  # Log the first 500 characters of the context
    
    # Perform question answering
    try:
        answer = qa_pipeline(question=question, context=context)
        app.logger.debug(f"Answer: {answer['answer']}")
        return jsonify({'answer': answer['answer']})
    except Exception as e:
        app.logger.error(f"Error during question answering: {e}")
        return jsonify({'answer': 'An error occurred while answering the question.'}), 500

@app.route('/generate_quiz', methods=['POST'])
def generate_quiz():
    data = request.json
    text = data.get('text', '')
    
    # Log the received text
    app.logger.debug(f"Received text for quiz generation: {text}")
    
    # Generate quiz questions
    try:
        questions = []
        chunks = split_text(text)
        for chunk in chunks:
            qg_input = f"generate questions: {chunk}"
            generated_questions = question_generator(qg_input, max_length=64, num_return_sequences=3)
            for q in generated_questions:
                question_text = q['generated_text']
                questions.append({
                    'questionText': question_text,
                    'answerOptions': [
                        {'answerText': 'Option 1', 'isCorrect': False},
                        {'answerText': 'Option 2', 'isCorrect': True},
                        {'answerText': 'Option 3', 'isCorrect': False},
                        {'answerText': 'Option 4', 'isCorrect': False},
                    ],
                })
        return jsonify({'questions': questions})
    except Exception as e:
        app.logger.error(f"Error during quiz generation: {e}")
        return jsonify({'error': 'An error occurred during quiz generation.'}), 500

if __name__ == '__main__':
    from transformers import pipeline
    question_generator = pipeline("text2text-generation", model="t5-small", from_pt=True)
    app.run(debug=True)

# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from transformers import pipeline
# import logging
# import os
# from werkzeug.utils import secure_filename
# from PyPDF2 import PdfReader
# import docx  # Correct import for python-docx

# app = Flask(__name__)
# CORS(app, resources={r"/*": {"origins": "*"}})  # Allow all origins

# # Initialize the summarization and question-answering pipelines
# summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
# qa_pipeline = pipeline("question-answering")

# # Configure logging
# logging.basicConfig(level=logging.DEBUG)

# UPLOAD_FOLDER = 'uploads'
# ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_file(filepath):
#     ext = filepath.rsplit('.', 1)[1].lower()
#     text = ""
#     try:
#         if ext == 'txt':
#             with open(filepath, 'r', encoding='utf-8') as file:
#                 text = file.read()
#         elif ext == 'pdf':
#             with open(filepath, 'rb') as file:
#                 reader = PdfReader(file)
#                 for page in reader.pages:
#                     text += page.extract_text()
#         elif ext == 'docx':
#             doc = docx.Document(filepath)
#             for para in doc.paragraphs:
#                 text += para.text
#         elif ext == 'doc':
#             # Handle .doc files if needed
#             pass
#     except Exception as e:
#         app.logger.error(f"Error extracting text from file: {e}")
#     return text

# @app.route('/summarize', methods=['POST'])
# def summarize_text():
#     if 'file' in request.files:
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400
#         if file and allowed_file(file.filename):
#             filename = secure_filename(file.filename)
#             filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             file.save(filepath)
#             text = extract_text_from_file(filepath)
#             if not text:
#                 return jsonify({'error': 'Failed to extract text from file'}), 500
#             try:
#                 app.logger.debug(f"Extracted text: {text[:500]}...")  # Log the first 500 characters of the extracted text
#                 summary = summarizer(text, max_length=150, min_length=100, do_sample=False)
#                 summary_text = summary[0]['summary_text']
#                 app.logger.debug(f"Summary: {summary_text}")
#                 return jsonify({'summary': summary_text})
#             except Exception as e:
#                 app.logger.error(f"Error during summarization: {e}")
#                 return jsonify({'summary': 'An error occurred during summarization.'}), 500
#         return jsonify({'error': 'Invalid file type'}), 400
#     else:
#         data = request.json
#         text = data.get('text', '')
        
#         # Log the received text
#         app.logger.debug(f"Received text: {text}")
        
#         # Perform summarization
#         try:
#             summary = summarizer(text, max_length=150, min_length=100, do_sample=False)
#             summary_text = summary[0]['summary_text']
            
#             # Log the summary
#             app.logger.debug(f"Summary: {summary_text}")
            
#             return jsonify({'summary': summary_text})
#         except Exception as e:
#             app.logger.error(f"Error during summarization: {e}")
#             return jsonify({'summary': 'An error occurred during summarization.'}), 500

# @app.route('/question_answer', methods=['POST'])
# def question_answer():
#     data = request.json
#     context = data.get('context', '')
#     question = data.get('question', '')
    
#     # Log the received question and context
#     app.logger.debug(f"Received question: {question}")
#     app.logger.debug(f"Context: {context[:500]}...")  # Log the first 500 characters of the context
    
#     # Perform question answering
#     try:
#         answer = qa_pipeline(question=question, context=context)
#         app.logger.debug(f"Answer: {answer['answer']}")
#         return jsonify({'answer': answer['answer']})
#     except Exception as e:
#         app.logger.error(f"Error during question answering: {e}")
#         return jsonify({'answer': 'An error occurred while answering the question.'}), 500

# if __name__ == '__main__':
#     app.run(debug=True)